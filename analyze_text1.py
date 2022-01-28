import sys
import os
import PyPDF2
import pandas as pd
from main import conn
from docx import Document
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request

app = Flask(__name__)
app.secret_key = 'key'
app.config["UPLOAD_FOLDER"] = "/TextAnalysis"

documents = []

with open('document.py', mode='w') as f:
    fi = 'doc = [' \
         '"""' \
         'temporary var' \
         '"""]'
    f.write("{}".format(fi))
f.close()

sys.path.append('TextAnalysis')
client = conn()
dat = {
    'Overall Positive Score': [],
    'Overall Neutral Score': [],
    'Overall Negative Score': [],
    'Sentence': [],
    'Positive Score': [],
    'Neutral Score': [],
    'Negative Score': [],
    'Target Text': [],
    'Target Sentiment': [],
    'Target Score Positive': [],
    'Target Score Negative': [],
    'Assessment Text': [],
    'Assessment Sentiment': [],
    'Assessment Score Positive': [],
    'Assessment Score Negative': [],
    'Detected Language': [],
    'Confidence Score': [],
    'Recognized Text': [],
    'Category': [],
    'Confidence Score Entity': [],
    'Linked Entity Name': [],
    'Linked Entity URL': [],
    'Linked Entity Confidence Score': [],
    'Key Phrases': [],
    'Key Phrases Row': []
}
key_str = ""
linked_entites = ""
linked_entites1 = ""
named_entities = ""
overall_positive_sentiment_score = ""
overall_neutral_sentiment_score = ""
overall_negative_sentiment_score = ""
a = ""
overall_sentiment = pd.DataFrame()
sentiment = pd.DataFrame()
language = pd.DataFrame()
key_phrase = pd.DataFrame()
entity_recognition = pd.DataFrame()
entity_linking = pd.DataFrame()
key_phrase_row = pd.DataFrame()


def word_to_text(file_name):
    document = Document(file_name)
    docum = []
    i = 0
    for i in range(0, len(document.paragraphs)):
        if len(document.paragraphs[i].text) > 0:
            if len(document.paragraphs[i].text) != 1 and document.paragraphs[i].text != '':
                docum.append(document.paragraphs[i].text)
        i = i + 1
    global documents
    documents = docum


def pdf_to_text(file_name):
    pdf_read = PyPDF2.PdfFileReader(file_name)
    for page_num in range(pdf_read.numPages):
        page_obj = pdf_read.getPage(page_num)
        txt = page_obj.extractText()
        # space = ''.center(100, '-')

        # write_file.write('{}\n'.format(space))
        # write_file.write('Page {}\n'.format(page_num + 1))
        # write_file.write('{}\n'.format(space))
        with open('document.py', mode='w') as doc:
            docu = 'doc = [' \
                   '"""' \
                   '{}' \
                   '"""]'.format(txt)
            doc.write("{}".format(docu))
        from document import doc
        global documents
        documents = doc


def analysis():
    # Data cleaning from old
    dat = {
        'Overall Positive Score': [],
        'Overall Neutral Score': [],
        'Overall Negative Score': [],
        'Sentence': [],
        'Positive Score': [],
        'Neutral Score': [],
        'Negative Score': [],
        'Target Text': [],
        'Target Sentiment': [],
        'Target Score Positive': [],
        'Target Score Negative': [],
        'Assessment Text': [],
        'Assessment Sentiment': [],
        'Assessment Score Positive': [],
        'Assessment Score Negative': [],
        'Detected Language': [],
        'Confidence Score': [],
        'Recognized Text': [],
        'Category': [],
        'Confidence Score Entity': [],
        'Linked Entity Name': [],
        'Linked Entity URL': [],
        'Linked Entity Confidence Score': [],
        'Key Phrases': [],
        'Key Phrases Row': []
    }
    # Language
    lan = client.detect_language(documents=documents)
    # print(lan)
    # dat['Detected Language'].clear()
    for lan in lan:
        global a
        a = format(lan.primary_language.name)
        dat['Detected Language'].append(format(lan.primary_language.name))
        dat['Confidence Score'].append(format(lan.primary_language.confidence_score))

    # Key Phrases
    keys = client.extract_key_phrases(documents=documents)

    for key in keys:
        res = str(key.key_phrases)[1:-1]
        res = res.replace("'", "")
        for k in key.key_phrases:
            # print(k)
            dat['Key Phrases Row'].append(k)
        dat['Key Phrases'].append(res)
    res = str(dat['Key Phrases Row'])[1:-1]
    res = res.replace("'", "")
    global key_str
    key_str = res
    # Recognized Entity
    rec_ent = client.recognize_entities(documents=documents)
    # print(rec_ent)
    for ent in rec_ent:
        for entity in ent.entities:
            dat['Recognized Text'].append(format(entity.text))
            dat['Category'].append(format(entity.category))
            dat['Confidence Score Entity'].append(format(entity.confidence_score))
    res = str(dat['Recognized Text'])[1:-1]
    res = res.replace("'", "")
    global named_entities
    named_entities = res

    # Recognized Linked Entity
    rec_ent_link = client.recognize_linked_entities(documents=documents)
    # print(rec_ent_link)
    for ent in rec_ent_link:
        for enti in ent.entities:
            dat['Linked Entity Name'].append(enti.name)
            dat['Linked Entity URL'].append(enti.url)
            for entity in enti.matches:
                dat['Linked Entity Confidence Score'].append(entity.confidence_score)
                break
    res = str(dat['Linked Entity Name'])[1:-1]
    res = res.replace("'", "")

    global linked_entites
    linked_entites = res

    res1 = str(dat['Linked Entity URL'])[1:-1]
    res1 = res1.replace("'", "")
    global linked_entites1
    linked_entites1 = res1

    # Sentiment
    response = client.analyze_sentiment(
        documents=documents, language='en-US', show_opinion_mining=True
    )

    # positive_reviews = [doc for doc in response if doc.sentiment == 'postive']
    # mixed_reviews = [doc for doc in response if doc.sentiment == 'mixed']
    # negative_reviews = [doc for doc in response if doc.sentiment == 'negative']

    # print(len(positive_reviews), len(mixed_reviews), len(negative_reviews))

    for document in response:
        # print('Postive score : ',document['confidence_scores']['positive'])
        # print('Negative score : ', document['confidence_scores']['negative'])
        # print('Neutral score : ', document['confidence_scores']['neutral'])
        tes = "Positive: ", document['confidence_scores']['positive']
        st = ''.join(map(str, tes))

        global overall_positive_sentiment_score
        overall_positive_sentiment_score = st

        tes = "Negative: ", document['confidence_scores']['negative']
        st = ''.join(map(str, tes))
        global overall_negative_sentiment_score
        overall_negative_sentiment_score = st

        tes = "Neutral: ", document['confidence_scores']['neutral']
        st = ''.join(map(str, tes))
        global overall_neutral_sentiment_score
        overall_neutral_sentiment_score = st
        # print(overall_positive_sentiment_score,overall_neutral_sentiment_score,overall_negative_sentiment_score)
        dat['Overall Positive Score'].append(document['confidence_scores']['positive'])
        dat['Overall Neutral Score'].append(document['confidence_scores']['neutral'])
        dat['Overall Negative Score'].append(document['confidence_scores']['negative'])
        # print('Sentiment Anaylsis Outcome: {0}'.format(document.sentiment))
        # print('Overall score : postive={0:.2f}; neutra={1:.2f};negative={2:.2f}'.format(
        #     document.confidence_scores.positive,
        #     document.confidence_scores.neutral,
        #     document.confidence_scores.negative,
        # ))
        # print('-' * 75)

        # break down the anaysis by each sentence
        sentences = document.sentences
        for indx, sentence in enumerate(sentences):
            # print('Sentence #{0}'.format(indx + 1))
            # print('Sentence Text: {0}'.format(sentence.text))
            # print('Sentence score:positive = {0:.2f}; neutral= {1:.2f}; negative={2:.2f}'.format(
            #     sentence.confidence_scores.positive,
            #     sentence.confidence_scores.neutral,
            #     sentence.confidence_scores.negative,
            # ))
            dat['Sentence'].append(sentence.text)
            dat['Positive Score'].append(sentence.confidence_scores.positive)
            dat['Neutral Score'].append(sentence.confidence_scores.neutral)
            dat['Negative Score'].append(sentence.confidence_scores.negative)

            # opinion mining result
            # print(sentence)
            for mined_opinion in sentence.mined_opinions:
                target = mined_opinion.target
                # print(target)
                # print('\t "{0}" target text; {1}'.format(target.sentiment, target.text))
                # print(r'\t Target Score: \n\tPositive={0:.2f} \Negative={1:.2f}'.format(
                #     target.confidence_scores.positive,
                #     target.confidence_scores.negative,
                # ))
                dat['Target Text'].append(target.text)
                dat['Target Sentiment'].append(target.sentiment)
                dat['Target Score Positive'].append(target.confidence_scores.positive)
                dat['Target Score Negative'].append(target.confidence_scores.negative)

                for assessment in mined_opinion.assessments:
                    # print('\t {0} assessment text : {1}'.format(assessment.sentiment, assessment.text))
                    # print(r'\t Assessment Score : \n\t Positive = {0:.2f} \Negative = {1:.2f}'.format(
                    #     assessment.confidence_scores.positive,
                    #     assessment.confidence_scores.negative,
                    # ))
                    dat['Assessment Text'].append(assessment.text)
                    dat['Assessment Sentiment'].append(assessment.sentiment)
                    dat['Assessment Score Positive'].append(assessment.confidence_scores.positive)
                    dat['Assessment Score Negative'].append(assessment.confidence_scores.negative)

        df_overall_sentiment = pd.DataFrame(dat, columns=[
            'Overall Positive Score',
            'Overall Neutral Score',
            'Overall Negative Score',
        ])

        df_sentiment = pd.DataFrame(dat, columns=[
            'Sentence',
            'Positive Score',
            'Neutral Score',
            'Negative Score'
        ])
        # df1 = pd.DataFrame(dat, columns=[
        #     'Target Text',
        #     'Target Sentiment',
        #     'Target Score Postive',
        #     'Target Score Negative',
        # ])

        # df2 = pd.DataFrame(dat, columns=[
        #     'Assessment Text',
        #     'Assessment Sentiment',
        #     'Assessment Score Positive',
        #     'Assessment Score Negative',
        # ])

        df_language = pd.DataFrame(dat, columns=[
            'Detected Language',
            'Confidence Score'
        ])

        df_ent_rec = pd.DataFrame(dat, columns=[
            'Recognized Text',
            'Category',
            'Confidence Score Entity'
        ])

        df_ent_link = pd.DataFrame(dat, columns=[
            'Linked Entity Name',
            'Linked Entity URL',
            'Linked Entity Confidence Score'
        ])

        df_key_phrase = pd.DataFrame(dat, columns=['Key Phrases'])

        df_key_phrase_row = pd.DataFrame(dat, columns=['Key Phrases Row'])

        global sentiment
        sentiment = df_sentiment

        global language
        language = df_language

        global key_phrase
        key_phrase = df_key_phrase

        global entity_recognition
        entity_recognition = df_ent_rec

        global entity_linking
        entity_linking = df_ent_link

        global key_phrase_row
        key_phrase_row = df_key_phrase_row

        global overall_sentiment
        overall_sentiment = df_overall_sentiment


@app.route('/', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST' and 'text-upload' in request.form:
        doc = [request.form['text']]
        global documents
        documents = doc
        analysis()

    elif request.method == 'POST' and 'file-upload' in request.form:
        if request.files:
            file = request.files["file-type"]
            print(file.filename)

            os.makedirs(os.path.join(app.instance_path, 'files'), exist_ok=True)
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.instance_path, 'files', secure_filename(filename)))

            extension = file.filename.split('.')[-1]
            if extension == 'docx':
                # time.sleep(2)
                word_to_text('instance/files/{}'.format(filename))
            elif extension == 'pdf':
                # time.sleep(2)
                pdf_to_text('instance/files/{}'.format(filename))
            else:
                return render_template('upload_file.html', error="error")
        analysis()
    return render_template('upload_file.html',
                           det_lang=a,
                           overall_positive_sentiment_score=overall_positive_sentiment_score,
                           overall_neutral_sentiment_score=overall_neutral_sentiment_score,
                           overall_negative_sentiment_score=overall_negative_sentiment_score,
                           named_entities=named_entities,
                           linked_entites=linked_entites,
                           linked_entites1=linked_entites1,
                           key_str = key_str,
                           tables_sentiment=[overall_sentiment.to_html(classes='data', header="true", index=False)],
                           tables_languages=[language.to_html(classes='data', header="true", index=False)],
                           tables_key_phrase=[key_phrase.to_html(classes='data', header="true", index=False)],
                           tables_entity_recognition=[
                               entity_recognition.to_html(classes='data', header="true", index=False)],
                           tables_entity_linking=[entity_linking.to_html(classes='data', header="true", index=False)]
                           )


@app.route('/export', methods=['POST'])
def export():
    writer = pd.ExcelWriter(r'export_dataframe.xlsx',
                            engine='xlsxwriter')
    language.to_excel(writer, sheet_name='Sheet1')
    overall_sentiment.to_excel(writer, sheet_name='Sheet1', startcol=6)
    sentiment.to_excel(writer, sheet_name='Sheet1', startcol=6, startrow=5)
    # df1.to_excel(writer, sheet_name='Sheet1', startcol=8)
    # df2.to_excel(writer, sheet_name='Sheet2', startcol=14)
    entity_recognition.to_excel(writer, sheet_name='Sheet1', startcol=13)
    entity_linking.to_excel(writer, sheet_name='Sheet1', startcol=19)
    key_phrase_row.to_excel(writer, sheet_name='Sheet1', startcol=25)
    writer.save()
    return render_template('upload_file.html', ress="Your file has been exported",
                           det_lang=a,
                           overall_positive_sentiment_score=overall_positive_sentiment_score,
                           overall_neutral_sentiment_score=overall_neutral_sentiment_score,
                           overall_negative_sentiment_score=overall_negative_sentiment_score,
                           named_entities=named_entities,
                           linked_entites=linked_entites,
                           linked_entites1=linked_entites1,
                           key_str = key_str)

